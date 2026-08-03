#!/usr/bin/env python3
"""Build WPU SABLAe Portal User Manual DOCX from USER_MANUAL.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "USER_MANUAL.md"
IMAGES = ROOT / "images"
OUT_DOCX = ROOT / "WPU_SABLAe_User_Manual.docx"


def set_run_font(run, name: str = "Calibri", size: int | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={0: 22, 1: 16, 2: 13, 3: 12}.get(level, 12), bold=True)
        if level <= 1:
            run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x2E)
    return p


def add_paragraph_text(doc: Document, text: str, *, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_bullet(doc: Document, text: str, level: int = 0):
    style = "List Number" if False else "List Bullet"
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_image(doc: Document, rel_path: str, caption: str | None = None):
    name = Path(rel_path).name
    path = IMAGES / name
    if not path.exists():
        add_paragraph_text(doc, f"[Missing image: {name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=9, bold=False)
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", raw):
            i += 1
            continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, size=10, bold=(r_idx == 0))
            if r_idx == 0:
                try:
                    from docx.oxml import OxmlElement

                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "E8F0EC")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
                except Exception:
                    pass
    doc.add_paragraph()


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def build() -> Path:
    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("WPU SABLAe Portal")
    set_run_font(r, size=26, bold=True)
    r.font.color.rgb = RGBColor(0x0B, 0x3D, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("User Manual")
    set_run_font(r, size=18, bold=True)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("SAFE WPU Scheduling and LMS Portal\nWestern Philippines University\nVersion 1.0 — July 2026")
    set_run_font(r, size=11)
    r.italic = True

    logo = IMAGES / "logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.6))

    doc.add_page_break()

    i = 0
    # Skip the markdown title/logo block already rendered
    while i < len(lines):
        line = lines[i]
        if line.startswith("# WPU SABLAe") or line.startswith("**SAFE") or line.startswith("Western") or line.startswith("Version"):
            i += 1
            continue
        if line.startswith("![WPU Logo]"):
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        break

    pending_caption: str | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, strip_md_inline(stripped[2:]), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, strip_md_inline(stripped[3:]), 1)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, strip_md_inline(stripped[4:]), 2)
            i += 1
            continue

        img = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img:
            alt, src = img.group(1), img.group(2)
            # Look ahead for italic caption
            caption = None
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("*") and lines[i + 1].strip().endswith("*"):
                caption = strip_md_inline(lines[i + 1].strip())
                i += 1
            add_image(doc, src, caption or alt)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            add_paragraph_text(doc, strip_md_inline(stripped), italic=True)
            i += 1
            continue

        if stripped.startswith("> "):
            tip = strip_md_inline(stripped[2:])
            p = doc.add_paragraph()
            run = p.add_run(tip)
            set_run_font(run, size=10, bold=False)
            run.italic = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = parse_table_rows(lines, i)
            add_table(doc, [[strip_md_inline(c) for c in row] for row in rows])
            i = next_i
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            add_numbered(doc, strip_md_inline(m_num.group(2)))
            i += 1
            continue

        if stripped.startswith("- "):
            add_bullet(doc, strip_md_inline(stripped[2:]))
            i += 1
            continue

        if stripped.startswith("*End of"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(strip_md_inline(stripped.strip("*")))
            set_run_font(r, size=10)
            r.italic = True
            i += 1
            continue

        add_paragraph_text(doc, strip_md_inline(stripped))
        i += 1

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}")
